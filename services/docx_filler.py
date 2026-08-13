# ============================================================
#  services/docx_filler.py
#  .docx shablonni to'ldirish
#  - Times New Roman shriftini barcha matnlarga o'rnatadi
#  - Shrift o'lchamini (font size) shablondagi 14pt o'lchamda 100% bir xil saqlaydi
#  - Faqat FIO, YONALISH, OQUV_YILI larni QALIN (Bold) qiladi
# ============================================================

from docx import Document
from docx.shared import Pt


def fill_template(template_path: str, output_path: str, data: dict) -> None:
    """
    template_path : .docx shablon fayli yo'li
    output_path   : natija .docx fayli yo'li
    data          : {"FIO": "...", "YONALISH": "...", "OQUV_YILI": "...", "SANA": "..."}
    """
    doc = Document(template_path)

    # 1. Paragraflar
    for para in doc.paragraphs:
        _process_paragraph(para, data)

    # 2. Jadvallar ichidagi kataklar
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph(para, data)

    # 3. Header va Footer
    for section in doc.sections:
        for para in section.header.paragraphs:
            _process_paragraph(para, data)
        for para in section.footer.paragraphs:
            _process_paragraph(para, data)

    doc.save(output_path)


def _process_paragraph(para, data: dict) -> None:
    if not para.text:
        return

    # Asl shrift o'lchamini aniqlaymiz (masalan 14pt)
    original_size = None
    for r in para.runs:
        if r.font and r.font.size:
            original_size = r.font.size
            break
    if not original_size:
        original_size = Pt(14)

    # Shriftni Times New Roman va aniq o'lchamga o'tkazish
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = original_size

    # 1. Har bir run ichida replace qilish
    for run in para.runs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(value))
                if key in ["FIO", "YONALISH", "OQUV_YILI"]:
                    run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = original_size

    # 2. Agar qavslar XML ichida bo'lingan bo'lsa (cross-run)
    full_text = "".join(r.text for r in para.runs)
    has_unreplaced = any(f"{{{{{key}}}}}" in full_text for key in data.keys())

    if has_unreplaced:
        _replace_and_format_paragraph(para, data, original_size)


def _replace_and_format_paragraph(para, data: dict, target_size) -> None:
    """
    Paragraf matnini almashtirib, faqat FIO, YONALISH va OQUV_YILI ni BOLD qiladi.
    Barcha matnlarga Times New Roman va bir xil Target Size (14pt) beradi.
    """
    text = "".join(r.text for r in para.runs)

    fio_val = data.get("FIO", "")
    yon_val = data.get("YONALISH", "")
    oq_val  = data.get("OQUV_YILI", "")
    sana_val = data.get("SANA", "")

    text = text.replace("{{FIO}}", fio_val)
    text = text.replace("{{YONALISH}}", yon_val)
    text = text.replace("{{OQUV_YILI}}", oq_val)
    text = text.replace("{{SANA}}", sana_val)

    is_italic = any(r.italic for r in para.runs)
    p_element = para._p
    for r in para.runs:
        p_element.remove(r._r)

    bold_tokens = [fio_val, yon_val, oq_val]
    bold_tokens = [t for t in bold_tokens if t]

    current_text = text
    while current_text:
        first_pos = len(current_text)
        found_token = None

        for tok in bold_tokens:
            pos = current_text.find(tok)
            if pos != -1 and pos < first_pos:
                first_pos = pos
                found_token = tok

        if found_token is not None:
            if first_pos > 0:
                r_norm = para.add_run(current_text[:first_pos])
                r_norm.font.name = "Times New Roman"
                r_norm.font.size = target_size
                r_norm.font.bold = False
                if is_italic:
                    r_norm.font.italic = True

            r_bold = para.add_run(found_token)
            r_bold.font.name = "Times New Roman"
            r_bold.font.size = target_size
            r_bold.font.bold = True
            if is_italic:
                r_bold.font.italic = True

            current_text = current_text[first_pos + len(found_token):]
        else:
            r_rest = para.add_run(current_text)
            r_rest.font.name = "Times New Roman"
            r_rest.font.size = target_size
            r_rest.font.bold = False
            if is_italic:
                r_rest.font.italic = True
            break
