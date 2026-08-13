import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = Document()

# Sahifa chegaralari (margins)
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# 1. Header Table (3 ustunli jadval)
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# Ustun kengliklari
widths = [Inches(2.5), Inches(1.5), Inches(2.5)]
hdr_cells = table.rows[0].cells

for i, w in enumerate(widths):
    hdr_cells[i].width = w

# Chap matn (Uzbek)
p_left = hdr_cells[0].paragraphs[0]
p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_l = p_left.add_run(
    "O’ZBEKISTON RESPUBLIKASI\n"
    "QASHQADARYO VILOYATI\n"
    "“QARSHI TIBBIYOT TEXNIKUMI”\n"
    "NODAVLAT TA’LIM MUASSASASI"
)
run_l.font.name = "Times New Roman"
run_l.font.size = Pt(8.5)
run_l.font.bold = True

# O'rta (Logo)
p_mid = hdr_cells[1].paragraphs[0]
p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_m = p_mid.add_run("Qarshi tibbiyot\ntexnikumi")
run_m.font.name = "Times New Roman"
run_m.font.size = Pt(9)
run_m.font.bold = True
run_m.font.color.rgb = RGBColor(0, 51, 153)

# O'ng matn (Russian)
p_right = hdr_cells[2].paragraphs[0]
p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = p_right.add_run(
    "РЕСПУБЛИКА УЗБЕКИСТАН\n"
    "КАШКАДАРЬИНСКАЯ ОБЛАСТЬ\n"
    "НЕГОСУДАРСТВЕННОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ\n"
    "«КАРШИНСКИЙ МЕДИЦИНСКИЙ ТЕХНИКУМ»"
)
run_r.font.name = "Times New Roman"
run_r.font.size = Pt(8.5)
run_r.font.bold = True

# Jadval chegaralariga kulrang nuqtali chiziq qo'shish
tblPr = table._tbl.tblPr
borders = parse_xml(
    r'<w:tblBorders %s>'
    r'  <w:top w:val="dotted" w:sz="4" w:space="0" w:color="888888"/>'
    r'  <w:left w:val="dotted" w:sz="4" w:space="0" w:color="888888"/>'
    r'  <w:bottom w:val="dotted" w:sz="4" w:space="0" w:color="888888"/>'
    r'  <w:right w:val="dotted" w:sz="4" w:space="0" w:color="888888"/>'
    r'  <w:insideH w:val="dotted" w:sz="4" w:space="0" w:color="888888"/>'
    r'  <w:insideV w:val="dotted" w:sz="4" w:space="0" w:color="888888"/>'
    r'</w:tblBorders>' % nsdecls('w')
)
tblPr.append(borders)

# 2. Ajratuvchi chiziq
p_space = doc.add_paragraph()
p_space.paragraph_format.space_before = Pt(6)
p_space.paragraph_format.space_after = Pt(6)
p_border = parse_xml(
    r'<w:pBdr %s><w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/></w:pBdr>' % nsdecls('w')
)
p_space._p.get_or_add_pPr().append(p_border)

# 3. Shahar va Sana
table_meta = doc.add_table(rows=1, cols=2)
table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
m_cells = table_meta.rows[0].cells
m_cells[0].width = Inches(3.25)
m_cells[1].width = Inches(3.25)

p_city = m_cells[0].paragraphs[0]
r_city = p_city.add_run("Qarshi shahri")
r_city.font.name = "Times New Roman"
r_city.font.size = Pt(11)

p_sana = m_cells[1].paragraphs[0]
p_sana.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_sana = p_sana.add_run("{{SANA}} y.")
r_sana.font.name = "Times New Roman"
r_sana.font.size = Pt(11)

# 4. Bo'sh joy
p_gap1 = doc.add_paragraph()
p_gap1.paragraph_format.space_before = Pt(36)

# 5. MA'LUMOTNOMA Sarlavhasi
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(24)
r_title = p_title.add_run("MA’LUMOTNOMA")
r_title.font.name = "Times New Roman"
r_title.font.size = Pt(14)
r_title.font.bold = True

# 6. Body Paragraph 1
p_b1 = doc.add_paragraph()
p_b1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_b1 = p_b1.add_run("Ushbu  ma’lumotnoma  shuni  tasdiqlaydiki,  haqiqatdan  ham")
r_b1.font.name = "Times New Roman"
r_b1.font.size = Pt(11)

# 7. Body Paragraph 2 (Placeholders)
p_b2 = doc.add_paragraph()
p_b2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_b2.paragraph_format.space_before = Pt(12)
p_b2.paragraph_format.space_after = Pt(12)

r_fio = p_b2.add_run("{{FIO}} ")
r_fio.font.bold = True
r_fio.font.name = "Times New Roman"
r_fio.font.size = Pt(11)

r_oq = p_b2.add_run("{{OQUV_YILI}}-o‘quv yilida ")
r_oq.font.bold = True
r_oq.font.name = "Times New Roman"
r_oq.font.size = Pt(11)

r_yon = p_b2.add_run("{{YONALISH}} ")
r_yon.font.bold = True
r_yon.font.name = "Times New Roman"
r_yon.font.size = Pt(11)

r_rest = p_b2.add_run("yo‘nalishiga shartnoma asosida o‘qishga qabul qilindi. Talaba o‘qishni 2026-yil sentyabr oyidan boshlaydi.")
r_rest.font.name = "Times New Roman"
r_rest.font.size = Pt(11)

# 8. Note
p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_note.paragraph_format.space_before = Pt(12)
p_note.paragraph_format.space_after = Pt(60)
r_note = p_note.add_run("Ma’lumotnoma so‘ralgan joyga taqdim etish uchun berildi")
r_note.font.italic = True
r_note.font.name = "Times New Roman"
r_note.font.size = Pt(10.5)

# 9. Footer (Imzo va Pechat qismi)
table_foot = doc.add_table(rows=1, cols=2)
table_foot.alignment = WD_TABLE_ALIGNMENT.CENTER
f_cells = table_foot.rows[0].cells
f_cells[0].width = Inches(3.5)
f_cells[1].width = Inches(3.0)

p_fl = f_cells[0].paragraphs[0]
r_fl = p_fl.add_run("“Qarshi tibbiyot texnikumi”\nijrochi direktori:")
r_fl.font.bold = True
r_fl.font.name = "Times New Roman"
r_fl.font.size = Pt(11)

p_fr = f_cells[1].paragraphs[0]
p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_fr = p_fr.add_run("Sh.Raxmonov")
r_fr.font.bold = True
r_fr.font.name = "Times New Roman"
r_fr.font.size = Pt(11)

out_dir = r"c:\Users\user\Desktop\CHAT\docbot\templates"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "malumotnoma.docx")
doc.save(out_path)
print(f"Hujjat yaratildi: {out_path}")
